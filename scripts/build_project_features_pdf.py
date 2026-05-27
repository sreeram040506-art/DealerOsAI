from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "artifacts/Auto_Profit_Hub_Project_Features_Workflow.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
      run.font.name = "Calibri"
      if level == 1:
          run.font.color.rgb = RGBColor(46, 116, 181)
          run.font.size = Pt(16)
      elif level == 2:
          run.font.color.rgb = RGBColor(46, 116, 181)
          run.font.size = Pt(13)
      else:
          run.font.color.rgb = RGBColor(31, 77, 120)
          run.font.size = Pt(12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        p.add_run(item)


def add_feature_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Area", "What It Does", "Main Capabilities"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="1F4D78")
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    for area, purpose, capabilities in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], area, bold=True)
        set_cell_text(cells[1], purpose)
        set_cell_text(cells[2], capabilities)
    for row in table.rows:
        row.cells[0].width = Inches(1.45)
        row.cells[1].width = Inches(2.15)
        row.cells[2].width = Inches(2.9)
    doc.add_paragraph()


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color="1F4D78")
        set_cell_shading(cells[0], "F2F4F7")
        set_cell_text(cells[1], value)
        cells[0].width = Inches(1.85)
        cells[1].width = Inches(4.65)
    doc.add_paragraph()


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    body = p.add_run(text)
    body.font.size = Pt(10)
    body.font.name = "Calibri"
    doc.add_paragraph()


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Auto Profit Hub")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Project Working, Workflow, and Feature Reference")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(85, 85, 85)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Vehicle inventory, document registry, sales, customer, and profitability management system")
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()
    add_callout(
        doc,
        "Purpose",
        "Auto Profit Hub is built to help a dealership track vehicles from purchase through repair, advertising, sale, customer follow-up, reporting, and document compliance."
    )

    add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "The project is a full-stack vehicle inventory and dealership operations platform. "
        "It combines a React frontend, Express API, Prisma data layer, MongoDB storage, protected user authentication, multi-tenant dealership data separation, and document-processing tools for purchase and sale records."
    )
    add_bullets(doc, [
        "Centralizes inventory, sales, expenses, repairs, advertising, customer data, reports, and used vehicle records.",
        "Supports staff, manager, admin, and platform admin workflows with role-aware screens.",
        "Automates parts of the paperwork process by extracting vehicle and sale data from documents.",
        "Calculates profitability from sale price, purchase cost, repair cost, and operating inputs.",
        "Keeps customer records connected to sold vehicles, visitors, leads, and follow-up categories."
    ])

    add_heading(doc, "2. Main System Workflow")
    add_numbered(doc, [
        "User signs in and the app loads dealership-specific data.",
        "Vehicle is added manually or from a purchase/source document.",
        "Inventory record stores VIN, make, model, year, mileage, color, title number, purchase cost, seller details, and document availability.",
        "Repairs, advertising, and business expenses are recorded and can be linked to vehicles where applicable.",
        "Bill of Sale is uploaded or sale is entered manually when the vehicle is sold.",
        "Sale record is created, vehicle status changes to Sold, profit is calculated, and the Used Vehicle Record is updated with disposition data.",
        "Customer details from the sale are stored in Customers; users can also add visitors, leads, follow-ups, and manually linked vehicles.",
        "Dashboard, Sales, Reports, Cash Flow, and Team Analytics summarize performance."
    ])

    add_heading(doc, "3. Feature Map")
    add_feature_table(doc, [
        ("Authentication", "Protects dealership data and routes.", "Login, registration, JWT sessions, role-based protected routes."),
        ("Dashboard", "Shows the operational snapshot.", "Inventory status, sales performance, profit summary, recent activity."),
        ("Inventory", "Tracks every vehicle lifecycle.", "Add/edit/delete vehicles, statuses, purchase data, documents, repairs, sale action."),
        ("Sales", "Stores finalized sales.", "Sale table/cards, revenue, profit, bill-of-sale documents, edit/delete sale records."),
        ("Customers", "Manages buyers, visitors, and leads.", "Manual add, edit contact, categories, linked inventory vehicles, search, import from sales."),
        ("Used Forms", "Handles document-driven workflows.", "Bill of Sale upload, repair bill upload, used vehicle form generation."),
        ("Registry", "Maintains document compliance data.", "Used Vehicle Record entries, acquisition/disposition details, document updates."),
        ("Repairs", "Adds vehicle reconditioning cost.", "Parts/labor cost, shop, repair description, vehicle profit impact."),
        ("Advertising", "Tracks campaign spend.", "Campaigns, platforms, dates, spend, optional vehicle link."),
        ("Expenses", "Tracks operating costs.", "Business expense categories, date, notes, edit/delete."),
        ("Reports", "Exports and summarizes results.", "Revenue report, sales breakdown, profit reporting."),
        ("Team", "Admin visibility into staff activity.", "Users, roles, sales activity, performance analytics."),
        ("Settings", "Dealership profile management.", "Name, address, phone, email, logo/profile information."),
        ("Super Admin", "Platform-level management.", "Dealership oversight and platform admin functions."),
        ("AI Assistant", "Helps query operations.", "Ask about inventory, sales, and profit context from within the app.")
    ])

    add_heading(doc, "4. Inventory and Vehicle Management")
    doc.add_paragraph(
        "Inventory is the starting point for the operational workflow. Each vehicle record stores key identification, purchase, cost, status, and document information."
    )
    add_bullets(doc, [
        "Vehicle fields include VIN, make, model, year, mileage, color, purchase date, title number, and status.",
        "Purchase data includes seller name/address, purchase price, buyer fee, transport, inspection, registration, payment method, and total purchase cost.",
        "Vehicle status can move through Available, Reserved, Sold, or Returned depending on business activity.",
        "Documents are tracked with flags for generated Used Vehicle Record, source document, and Bill of Sale.",
        "Vehicle detail dialog provides tabs for overview, repairs, advertising, sale processing, documents, and customer viewing notes."
    ])

    add_heading(doc, "5. Sales and Profit Workflow")
    doc.add_paragraph(
        "Sales can be created manually from a vehicle detail screen or automatically when a Bill of Sale document is processed. "
        "The sale stores buyer information, sale date, sale price, payment method, and document archive data."
    )
    add_key_value_table(doc, [
        ("Profit formula", "Sale price - purchase cost - repair cost."),
        ("Sale action", "Marks the vehicle as Sold and creates or updates the sale record."),
        ("Document update", "Regenerates the Used Vehicle Record with disposition details."),
        ("Customer update", "Creates or updates a customer record with contact, category, and linked vehicle metadata."),
        ("Bill of Sale", "Can be uploaded, previewed, downloaded, and stored as base64 document data.")
    ])

    add_heading(doc, "6. Customers Section")
    doc.add_paragraph(
        "The Customers section stores people connected to the dealership. It supports buyers imported from sales and manual customers such as visitors, leads, and follow-ups."
    )
    add_bullets(doc, [
        "Manual customer creation with first name, last name, phone, email, address, city, state, and zip.",
        "Customer categories: Bought Vehicle, Came for Visit, Lead, Follow Up, and Other.",
        "Vehicle linking from Inventory with search by VIN, make, model, year, or status.",
        "Main customer search across name, contact number, email, category, address, and linked vehicle.",
        "Sales import creates or updates customer records from existing sale data.",
        "Existing customers can be edited to update contact details, category, and vehicle link."
    ])

    add_heading(doc, "7. Document Processing and Registry")
    doc.add_paragraph(
        "The system includes document workflows for dealership paperwork. Purchase/source documents and Bills of Sale can be processed to extract vehicle and transaction details."
    )
    add_bullets(doc, [
        "Document parser extracts VIN, year, make, model, title information, seller/source details, buyer/disposition details, dates, and prices.",
        "Used Vehicle Record PDF generation fills acquisition and disposition sections.",
        "Registry keeps VIN-based document records with source file name, document type, purchase details, and sale/disposition details.",
        "Document download and preview functions let users access records without manually searching file storage.",
        "Bulk import tooling can align purchase and sale documents to inventory records."
    ])

    add_heading(doc, "8. Financial and Operational Modules")
    add_feature_table(doc, [
        ("Repairs", "Captures vehicle-specific reconditioning costs.", "Repair shop, parts cost, labor cost, repair date, description."),
        ("Advertising", "Tracks marketing spend and optional vehicle association.", "Campaign name, platform, start/end dates, amount spent, linked vehicle."),
        ("Business Expenses", "Captures non-vehicle operating expenses.", "Category, amount, date, notes."),
        ("Cash Flow", "Summarizes money movement.", "Income and cost visibility for managers/admins."),
        ("Reports", "Provides management exports and summaries.", "Revenue reports, unit counts, gross revenue, net profit, buyer breakdowns.")
    ])

    add_heading(doc, "9. User Roles and Access")
    add_key_value_table(doc, [
        ("Staff", "Operational access to inventory, sales, used forms, customers, and day-to-day records. Profit can be masked in some views."),
        ("Manager", "Broader access to registry, reports, customers, and sales management functions."),
        ("Admin", "Full dealership management including team analytics, settings, reports, customers, inventory, and operations."),
        ("Super Admin", "Platform-level administration across dealerships.")
    ])

    add_heading(doc, "10. Technical Architecture")
    add_key_value_table(doc, [
        ("Frontend", "React, TypeScript, Vite, React Router, TanStack Query, Tailwind/shadcn-style UI components."),
        ("Backend", "Node.js, Express routes, authentication middleware, tenant injection middleware, file upload handling."),
        ("Database", "MongoDB through Prisma Client with models for dealership, users, vehicles, sales, purchases, repairs, customers, expenses, advertising, and registry."),
        ("Document services", "Document parser and used vehicle PDF service for OCR/AI extraction and generated paperwork."),
        ("Caching", "In-memory route cache for selected list endpoints such as vehicles and sales."),
        ("Security", "JWT authentication, role-aware protected routes, helmet, CORS, rate limiting, and dealership isolation.")
    ])

    add_heading(doc, "11. Key Data Models")
    add_feature_table(doc, [
        ("Dealership", "Tenant container.", "Users, vehicles, sales, purchases, repairs, customers, documents, expenses."),
        ("User", "Authenticated app user.", "Email, password, name, role, dealership relation."),
        ("Vehicle", "Inventory item.", "VIN, make/model/year, mileage, status, title number, purchase, sale, repairs."),
        ("Purchase", "Acquisition record.", "Seller, purchase price, fees, total cost, document data."),
        ("Sale", "Disposition record.", "Customer name, phone, address, sale date, sale price, payment, profit, Bill of Sale."),
        ("Customer", "Buyer/visitor/lead.", "Name, email, phone, address, license, notes metadata."),
        ("DocumentRegistry", "Compliance document record.", "Vehicle identity, acquisition and disposition fields, document base64."),
        ("Repair", "Reconditioning record.", "Vehicle, shop, parts, labor, date, description."),
        ("Expense/Advertising", "Cost tracking.", "Business cost and campaign spend records.")
    ])

    add_heading(doc, "12. Recent Customer Enhancements")
    add_bullets(doc, [
        "Customers route is mounted under /api/customers.",
        "Customers page is available in desktop sidebar and mobile navigation.",
        "Sales and Bill of Sale flows create or update customers automatically.",
        "Manual Add Customer workflow supports categories and inventory vehicle linking.",
        "Customer import from sales updates existing customer records instead of only creating new ones.",
        "Customer categories and vehicle links are stored in the existing notes field to avoid requiring an immediate database migration."
    ])

    add_heading(doc, "13. Operational Notes")
    add_bullets(doc, [
        "Restart the backend server after backend route changes.",
        "Run npm run build after frontend changes to verify TypeScript and Vite production build.",
        "If Prisma schema changes are made, regenerate Prisma Client before running the server.",
        "Avoid pushing local changes to GitHub unless explicitly requested.",
        "Keep customer contact data accurate because it is reused across sales records and customer follow-up workflows."
    ])

    add_heading(doc, "14. End-to-End Example")
    add_numbered(doc, [
        "Add a vehicle to Inventory from purchase details or source document.",
        "Attach repair bills and advertising spend as work is completed.",
        "Upload a Bill of Sale or manually process the sale from the vehicle record.",
        "The system marks the vehicle Sold, calculates profit, updates the Used Vehicle Record, and saves the Bill of Sale.",
        "The buyer appears in Customers as Bought Vehicle and can be edited with email, phone, category, and vehicle link.",
        "Managers review Sales, Reports, Cash Flow, and Dashboard totals."
    ])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Auto Profit Hub - Project Features and Workflow Reference")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
